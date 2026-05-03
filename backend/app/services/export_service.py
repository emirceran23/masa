"""Export service — generate a revised contract DOCX by applying accepted revisions.

For each clause in the contract:
  - If the clause has an 'accepted' revision, use revision.suggested_text.
  - If the clause has an 'edited' revision, use revision.edited_text.
  - Otherwise, use clause.original_text.

The resulting DOCX is uploaded to MinIO and the URL is returned.
"""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.exceptions import NotFoundError
from app.models.clause import Clause
from app.models.contract import Contract
from app.storage.minio_client import upload_file


async def export_revised_contract(
    db: AsyncSession,
    contract_id: uuid.UUID,
    user_id: uuid.UUID,
) -> tuple[bytes, str]:
    """Build a DOCX with accepted/edited revisions applied.

    Returns (file_bytes, storage_path).
    """
    contract = await _get_owned_contract(db, contract_id, user_id)

    # Load clauses with revisions
    stmt = (
        select(Clause)
        .options(selectinload(Clause.revisions))
        .where(Clause.contract_id == contract_id)
        .order_by(Clause.sequence_no.asc())
    )
    clauses = (await db.execute(stmt)).scalars().all()

    file_bytes = _build_docx(contract, clauses)

    object_name = (
        f"exports/{contract_id}/revised_"
        f"{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.docx"
    )
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    storage_path = upload_file(object_name, file_bytes, content_type)

    return file_bytes, storage_path


def _pick_final_text(clause: Clause) -> tuple[str, bool]:
    """Return (final_text, was_revised).

    Priority: edited > accepted > original.
    Scans all revisions so the highest-priority status wins.
    """
    edited_text: str | None = None
    accepted_text: str | None = None

    for rev in clause.revisions:
        if rev.status == "edited" and rev.edited_text:
            edited_text = rev.edited_text  # latest edited wins
        elif rev.status == "accepted" and accepted_text is None:
            accepted_text = rev.suggested_text

    if edited_text is not None:
        return edited_text, True
    if accepted_text is not None:
        return accepted_text, True
    return clause.original_text, False


def _build_docx(contract: Contract, clauses: list[Clause]) -> bytes:
    doc = Document()

    # Document title
    title = doc.add_heading(contract.file_name, level=0)
    title.alignment = 1  # CENTER

    doc.add_paragraph(
        f"Revize Edilmiş Sözleşme — Oluşturulma: "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d')}"
    )
    doc.add_paragraph()

    for clause in clauses:
        final_text, was_revised = _pick_final_text(clause)

        # Clause header
        header_p = doc.add_paragraph()
        run = header_p.add_run(
            f"Madde {clause.sequence_no}"
            + (f" ({clause.category})" if clause.category else "")
        )
        run.bold = True
        run.font.size = Pt(11)
        if was_revised:
            run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)  # green — revised

        # Clause body
        body_p = doc.add_paragraph(final_text)
        body_p.paragraph_format.space_after = Pt(8)

        # Revision note (subtle)
        if was_revised:
            note_p = doc.add_paragraph()
            note_run = note_p.add_run("↑ Bu madde revize edilmiştir.")
            note_run.italic = True
            note_run.font.size = Pt(8)
            note_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _get_owned_contract(
    db: AsyncSession,
    contract_id: uuid.UUID,
    user_id: uuid.UUID,
) -> Contract:
    stmt = select(Contract).where(
        Contract.id == contract_id, Contract.user_id == user_id
    )
    contract = (await db.execute(stmt)).scalar_one_or_none()
    if not contract:
        raise NotFoundError("Sözleşme bulunamadı.")
    return contract
