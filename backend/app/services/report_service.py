"""Report service — aggregates contract data and generates PDF / DOCX reports.

Two report types:
  summary  — overview stats: risk distribution, clause statuses, missing provisions count
  detailed — all of the above + per-clause breakdown with risk scores and revisions
"""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone
from typing import Any

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.exceptions import BadRequestError, NotFoundError
from app.models.clause import Clause
from app.models.contract import Contract
from app.models.report import Report
from app.storage.minio_client import upload_file

VALID_REPORT_TYPES = {"summary", "detailed"}
VALID_FORMATS = {"pdf", "docx"}

# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────


async def generate_report(
    db: AsyncSession,
    contract_id: uuid.UUID,
    user_id: uuid.UUID,
    report_type: str = "summary",
    fmt: str = "pdf",
) -> Report:
    """Generate a report, upload to MinIO, persist a Report row, and return it."""
    if report_type not in VALID_REPORT_TYPES:
        raise BadRequestError(f"Geçersiz rapor türü: {report_type}")
    if fmt not in VALID_FORMATS:
        raise BadRequestError(f"Geçersiz format: {fmt}")

    contract = await _get_owned_contract(db, contract_id, user_id)
    if contract.status != "analyzed":
        raise BadRequestError("Rapor üretmek için sözleşmenin analiz edilmiş olması gerekir.")

    payload = await _build_payload(db, contract)

    if fmt == "pdf":
        file_bytes = _render_pdf(contract, payload, report_type)
        content_type = "application/pdf"
        ext = "pdf"
    else:
        file_bytes = _render_docx(contract, payload, report_type)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"

    object_name = f"reports/{contract_id}/{report_type}_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.{ext}"
    storage_path = upload_file(object_name, file_bytes, content_type)

    report = Report(
        contract_id=contract_id,
        report_type=report_type,
        total_clauses=payload["total_clauses"],
        summary_data=payload["summary"],
        storage_path=storage_path,
    )
    db.add(report)
    await db.flush()
    await db.refresh(report)
    return report


async def list_reports(
    db: AsyncSession,
    contract_id: uuid.UUID,
    user_id: uuid.UUID,
) -> list[Report]:
    await _get_owned_contract(db, contract_id, user_id)
    stmt = (
        select(Report)
        .where(Report.contract_id == contract_id)
        .order_by(Report.created_at.desc())
    )
    rows = (await db.execute(stmt)).scalars().all()
    return list(rows)


async def get_report(
    db: AsyncSession,
    report_id: uuid.UUID,
    user_id: uuid.UUID,
) -> Report:
    stmt = (
        select(Report)
        .join(Contract, Contract.id == Report.contract_id)
        .where(Report.id == report_id, Contract.user_id == user_id)
    )
    report = (await db.execute(stmt)).scalar_one_or_none()
    if not report:
        raise NotFoundError("Rapor bulunamadı.")
    return report


# ─────────────────────────────────────────────────────────────────────────────
# Data aggregation
# ─────────────────────────────────────────────────────────────────────────────


async def _build_payload(db: AsyncSession, contract: Contract) -> dict[str, Any]:
    """Fetch all related data and return a structured dict used by both renderers."""
    # Clauses with risk assessments and revisions eagerly loaded
    stmt = (
        select(Clause)
        .options(
            selectinload(Clause.risk_assessment),
            selectinload(Clause.revisions),
        )
        .where(Clause.contract_id == contract.id)
        .order_by(Clause.sequence_no.asc())
    )
    clauses = (await db.execute(stmt)).scalars().all()

    # Missing provisions (stored in raw SQL table)
    missing_rows = await db.execute(
        text(
            "SELECT id, playbook_rule_id, description FROM missing_provisions "
            "WHERE contract_id = :cid ORDER BY detected_at ASC"
        ),
        {"cid": contract.id},
    )
    missing_provisions = [
        {"key": str(r.playbook_rule_id) if r.playbook_rule_id else "unknown", "description": r.description}
        for r in missing_rows
    ]

    # Risk distribution counts
    risk_counts: dict[str, int] = {"low": 0, "medium": 0, "high": 0, "unassessed": 0}
    status_counts: dict[str, int] = {}
    revision_counts: dict[str, int] = {"accepted": 0, "rejected": 0, "pending": 0, "edited": 0}

    clause_details: list[dict[str, Any]] = []

    for clause in clauses:
        ra = clause.risk_assessment
        level = ra.risk_level if ra else "unassessed"
        risk_counts[level] = risk_counts.get(level, 0) + 1

        status_counts[clause.status] = status_counts.get(clause.status, 0) + 1

        for rev in clause.revisions:
            revision_counts[rev.status] = revision_counts.get(rev.status, 0) + 1

        # Per-clause detail (used only in detailed report)
        clause_details.append(
            {
                "sequence_no": clause.sequence_no,
                "category": clause.category or "—",
                "status": clause.status,
                "original_text": clause.original_text,
                "risk_level": level,
                "commercial_score": ra.commercial_score if ra else None,
                "legal_score": ra.legal_score if ra else None,
                "rationale": ra.rationale if ra else None,
                "policy_compliance": ra.policy_compliance if ra else None,
                "cross_validated": ra.cross_validated if ra else False,
                "revisions": [
                    {
                        "suggested_text": rev.suggested_text,
                        "status": rev.status,
                        "edited_text": rev.edited_text,
                    }
                    for rev in clause.revisions
                ],
            }
        )

    summary: dict[str, Any] = {
        "risk_counts": risk_counts,
        "status_counts": status_counts,
        "revision_counts": revision_counts,
        "missing_provisions_count": len(missing_provisions),
        "missing_provisions": missing_provisions,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "contract_name": contract.file_name,
    }

    return {
        "total_clauses": len(clauses),
        "summary": summary,
        "clause_details": clause_details,
    }


# ─────────────────────────────────────────────────────────────────────────────
# PDF renderer (fpdf2)
# ─────────────────────────────────────────────────────────────────────────────

_RISK_COLORS = {
    "high":       (220, 38, 38),
    "medium":     (234, 88, 12),
    "low":        (22, 163, 74),
    "unassessed": (107, 114, 128),
}

# fpdf2 core fonts cover Latin-1 (cp1252) only.  Map Turkish chars that fall
# outside cp1252 so PDF rendering never raises a UnicodeEncodeError.
_TR_PDF = str.maketrans({
    "ğ": "g", "Ğ": "G",
    "ı": "i", "İ": "I",
    "ş": "s", "Ş": "S",   # ş/Ş are at 0x9A/0x8A in cp1252 on some systems;
    # keep them mapped just to be safe with strict Latin-1 environments
})


def _p(s: str) -> str:
    """Sanitise text for fpdf2 core-font rendering (Latin-1 safe)."""
    return s.translate(_TR_PDF) if s else ""

_STATUS_TR = {
    "draft":     "Taslak",
    "in_review": "İncelemede",
    "approved":  "Onaylandı",
    "rejected":  "Reddedildi",
}

_RISK_TR = {
    "high":       "Yüksek",
    "medium":     "Orta",
    "low":        "Düşük",
    "unassessed": "Değerlendirilmedi",
}


class _PDF(FPDF):
    """Custom FPDF subclass with branded header/footer."""

    def header(self):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(30, 64, 175)  # indigo-700
        self.cell(0, 8, "Masa - Hukuki Sozlesme Analizi", align="L")
        self.set_text_color(0, 0, 0)
        self.ln(4)
        self.set_draw_color(209, 213, 219)
        self.set_line_width(0.3)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(107, 114, 128)
        self.cell(0, 10, f"Sayfa {self.page_no()}", align="C")  # noqa: ascii-only ok


def _render_pdf(contract: Contract, payload: dict[str, Any], report_type: str) -> bytes:
    summary = payload["summary"]
    risk_counts = summary["risk_counts"]
    status_counts = summary["status_counts"]
    total = payload["total_clauses"]

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # ── Cover / title ──
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(17, 24, 39)
    pdf.multi_cell(0, 10, _p(contract.file_name), align="L")
    pdf.ln(2)

    report_label = "Ozet Rapor" if report_type == "summary" else "Detayli Rapor"
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 7, f"{report_label}  |  Olusturulma: {summary['generated_at'][:10]}", align="L")
    pdf.ln(8)

    # ── Risk distribution ──
    _pdf_section_title(pdf, "Risk Dagilimi")
    _pdf_risk_bar(pdf, risk_counts, total)
    pdf.ln(4)

    _pdf_stat_row(pdf, [
        ("Toplam Madde", str(total)),
        ("Yuksek Risk", str(risk_counts.get("high", 0))),
        ("Orta Risk",   str(risk_counts.get("medium", 0))),
        ("Dusuk Risk",  str(risk_counts.get("low", 0))),
    ])
    pdf.ln(6)

    # ── Clause status breakdown ──
    _pdf_section_title(pdf, "Madde Durumlari")
    for status, count in status_counts.items():
        label = _STATUS_TR.get(status, status)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(55, 65, 81)
        pdf.cell(60, 6, f"  {_p(label)}:", align="L")
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, str(count), align="L")
        pdf.ln()
    pdf.ln(4)

    # ── Missing provisions ──
    missing = summary["missing_provisions"]
    _pdf_section_title(pdf, f"Eksik Hukumler ({len(missing)})")
    if missing:
        for mp in missing:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(180, 83, 9)
            pdf.cell(0, 5, _p(f"  * {mp['key']}"), align="L")
            pdf.ln()
            if mp.get("description"):
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(55, 65, 81)
                pdf.multi_cell(0, 5, _p(f"    {mp['description']}"), align="L")
            pdf.ln(1)
    else:
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(107, 114, 128)
        pdf.cell(0, 6, "  Eksik hukum tespit edilmedi.", align="L")
        pdf.ln()
    pdf.ln(4)

    # ── Detailed clause breakdown ──
    if report_type == "detailed":
        _pdf_section_title(pdf, "Madde Detaylari")
        for cd in payload["clause_details"]:
            _pdf_clause_block(pdf, cd)

    return bytes(pdf.output())


def _pdf_section_title(pdf: _PDF, title: str) -> None:
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(0, 8, _p(title), align="L")
    pdf.ln(1)
    pdf.set_draw_color(199, 210, 254)
    pdf.set_line_width(0.4)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)


def _pdf_risk_bar(pdf: _PDF, risk_counts: dict, total: int) -> None:
    """Render a colored horizontal bar proportional to risk level counts."""
    if total == 0:
        return
    bar_w = pdf.w - pdf.l_margin - pdf.r_margin
    bar_h = 6
    x_start = pdf.l_margin
    y = pdf.get_y()
    for level in ("high", "medium", "low", "unassessed"):
        count = risk_counts.get(level, 0)
        if count == 0:
            continue
        segment_w = bar_w * count / total
        r, g, b = _RISK_COLORS.get(level, (107, 114, 128))
        pdf.set_fill_color(r, g, b)
        pdf.rect(x_start, y, segment_w, bar_h, "F")
        x_start += segment_w
    pdf.ln(bar_h + 3)

    # Legend
    pdf.set_font("Helvetica", "", 8)
    for level, (r, g, b) in _RISK_COLORS.items():
        count = risk_counts.get(level, 0)
        if count == 0:
            continue
        pdf.set_fill_color(r, g, b)
        pdf.rect(pdf.get_x(), pdf.get_y() + 1.5, 4, 4, "F")
        pdf.set_x(pdf.get_x() + 5)
        pdf.set_text_color(55, 65, 81)
        pdf.cell(35, 6, _p(f"{_RISK_TR[level]}: {count}"), align="L")
    pdf.ln()


def _pdf_stat_row(pdf: _PDF, stats: list[tuple[str, str]]) -> None:
    n = len(stats)
    col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / n
    y_top = pdf.get_y()

    # Row 1: large values
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(17, 24, 39)
    for i, (_, value) in enumerate(stats):
        x = pdf.l_margin + i * col_w
        pdf.set_fill_color(243, 244, 246)
        pdf.rect(x, y_top, col_w - 2, 14, "F")
        pdf.set_xy(x, y_top)
        pdf.cell(col_w - 2, 8, value, align="C")

    # Row 2: small labels
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(107, 114, 128)
    for i, (label, _) in enumerate(stats):
        x = pdf.l_margin + i * col_w
        pdf.set_xy(x, y_top + 8)
        pdf.cell(col_w - 2, 6, _p(label), align="C")

    pdf.set_xy(pdf.l_margin, y_top + 14)
    pdf.ln(0)


def _pdf_clause_block(pdf: _PDF, cd: dict) -> None:
    level = cd["risk_level"]
    r, g, b = _RISK_COLORS.get(level, (107, 114, 128))

    # Colored left border indicator via filled rect
    pdf.set_fill_color(r, g, b)
    pdf.rect(pdf.l_margin, pdf.get_y(), 2, 5, "F")

    pdf.set_x(pdf.l_margin + 4)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(17, 24, 39)
    pdf.cell(
        0, 5,
        _p(f"Madde #{cd['sequence_no']}  -  {cd['category']}  -  {_RISK_TR.get(level, level)}"),
        align="L",
    )
    pdf.ln(6)

    # Original text (truncated for summary)
    original = cd["original_text"]
    if len(original) > 400:
        original = original[:400] + "..."
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(55, 65, 81)
    pdf.multi_cell(0, 4.5, _p(original), align="L")
    pdf.ln(2)

    # Risk scores
    if cd["commercial_score"] is not None:
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(107, 114, 128)
        pdf.cell(
            0, 4,
            _p(
                f"  Ticari: {cd['commercial_score']*100:.0f}%  |  "
                f"Hukuki: {cd['legal_score']*100:.0f}%  |  "
                f"Politika Uyumu: {'Evet' if cd['policy_compliance'] else 'Hayir'}"
            ),
            align="L",
        )
        pdf.ln(4)

    # Rationale
    if cd.get("rationale"):
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(75, 85, 99)
        pdf.multi_cell(0, 4, _p(f"  {cd['rationale']}"), align="L")
        pdf.ln(1)

    # Revisions
    for rev in cd["revisions"]:
        status_label = {
            "accepted": "[Kabul]",
            "rejected": "[Red]",
            "pending":  "[Bekliyor]",
            "edited":   "[Duzenlendi]",
        }.get(rev["status"], f"[{rev['status']}]")
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(107, 114, 128)
        suggested = (rev["edited_text"] or rev["suggested_text"])[:200]
        pdf.multi_cell(0, 4, _p(f"  {status_label} {suggested}"), align="L")
        pdf.ln(1)

    pdf.ln(3)
    pdf.set_draw_color(229, 231, 235)
    pdf.set_line_width(0.2)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(3)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX renderer (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

_RISK_RGB = {
    "high":       RGBColor(0xDC, 0x26, 0x26),
    "medium":     RGBColor(0xEA, 0x58, 0x0C),
    "low":        RGBColor(0x16, 0xA3, 0x4A),
    "unassessed": RGBColor(0x6B, 0x72, 0x80),
}


def _render_docx(contract: Contract, payload: dict[str, Any], report_type: str) -> bytes:
    summary = payload["summary"]
    risk_counts = summary["risk_counts"]
    total = payload["total_clauses"]

    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(contract.file_name)
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph(
        f"{'Özet Rapor' if report_type == 'summary' else 'Detaylı Rapor'} | "
        f"Oluşturulma: {summary['generated_at'][:10]}"
    ).runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # ── Risk distribution ──
    h = doc.add_heading("Risk Dağılımı", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["Risk", "Madde Sayısı", "Yüzde", ""]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True

    for level in ("high", "medium", "low", "unassessed"):
        count = risk_counts.get(level, 0)
        row = tbl.add_row().cells
        row[0].text = _RISK_TR.get(level, level)
        row[1].text = str(count)
        row[2].text = f"{count/total*100:.1f}%" if total else "—"
        row[3].text = ""
        run = row[0].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = _RISK_RGB.get(level, RGBColor(0, 0, 0))

    doc.add_paragraph()

    # ── Status breakdown ──
    h2 = doc.add_heading("Madde Durumları", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    for status, count in summary["status_counts"].items():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{_STATUS_TR.get(status, status)}: ").bold = True
        p.add_run(str(count))

    doc.add_paragraph()

    # ── Missing provisions ──
    missing = summary["missing_provisions"]
    h3 = doc.add_heading(f"Eksik Hükümler ({len(missing)})", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    if missing:
        for mp in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(mp["key"] + ": ").bold = True
            p.add_run(mp.get("description") or "")
    else:
        doc.add_paragraph("Eksik hüküm tespit edilmedi.").runs[0].font.italic = True

    doc.add_paragraph()

    # ── Detailed clause breakdown ──
    if report_type == "detailed":
        h4 = doc.add_heading("Madde Detayları", level=1)
        h4.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        for cd in payload["clause_details"]:
            level = cd["risk_level"]

            # Clause heading
            clause_h = doc.add_heading(
                f"Madde #{cd['sequence_no']} — {cd['category']} — {_RISK_TR.get(level, level)}",
                level=2,
            )
            if clause_h.runs:
                clause_h.runs[0].font.color.rgb = _RISK_RGB.get(level, RGBColor(0x1E, 0x40, 0xAF))

            # Original text
            doc.add_paragraph(cd["original_text"])

            # Risk scores
            if cd["commercial_score"] is not None:
                info = doc.add_paragraph()
                info.add_run(
                    f"Ticari: {cd['commercial_score']*100:.0f}%  |  "
                    f"Hukuki: {cd['legal_score']*100:.0f}%  |  "
                    f"Politika Uyumu: {'Evet' if cd['policy_compliance'] else 'Hayır'}"
                )
                if info.runs:
                    info.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                    info.runs[0].font.size = Pt(9)

            if cd.get("rationale"):
                rat_p = doc.add_paragraph(cd["rationale"])
                if rat_p.runs:
                    rat_p.runs[0].italic = True
                    rat_p.runs[0].font.size = Pt(9)

            # Revisions
            for rev in cd["revisions"]:
                status_label = {
                    "accepted": "Kabul",
                    "rejected": "Red",
                    "pending": "Bekliyor",
                    "edited": "Düzenlendi",
                }.get(rev["status"], rev["status"])
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{status_label}] ").bold = True
                p.add_run(rev["edited_text"] or rev["suggested_text"])
                if p.runs:
                    p.runs[-1].font.size = Pt(9)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────


async def _get_owned_contract(
    db: AsyncSession,
    contract_id: uuid.UUID,
    user_id: uuid.UUID,
) -> Contract:
    stmt = select(Contract).where(
        Contract.id == contract_id, Contract.user_id == user_id
    )
    contract = (await db.execute(stmt)).scalar_one_or_none()
    if not contract:
        raise NotFoundError("Sözleşme bulunamadı.")
    return contract
